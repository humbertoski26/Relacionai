"""
Extracción de texto desde los formatos que puede subir una persona:
texto plano (.txt), Word (.docx) y PDF (.pdf).

Cada función lanza ExtractError con un mensaje pensado para mostrarse
directo a la persona que está subiendo el relato (en español, sin jerga
técnica).
"""

import io
import time

import pypdf
from docx import Document


class ExtractError(Exception):
    pass


MAX_CHARS = 120_000  # tope de caracteres que guardamos por relato
MAX_PDF_PAGINAS = 300  # tope de páginas a procesar, para no consumir memoria de más con PDFs enormes
MAX_PDF_SEGUNDOS = 25  # tope de tiempo acumulado leyendo páginas — algunos PDF (estructuras
# raras, casi escaneados) hacen que pypdf tarde muchísimo página por página; mejor cortar y
# devolver lo que se alcanzó a leer que quedarse pegado (esto se procesa en segundo plano,
# así que no bloquea al encargado, pero igual conviene no dejarlo corriendo para siempre)


def extraer_texto(nombre_archivo: str, contenido_bytes: bytes) -> str:
    nombre = (nombre_archivo or "").lower()

    if nombre.endswith(".txt"):
        texto = _decode_texto(contenido_bytes)
    elif nombre.endswith(".docx"):
        texto = _extraer_docx(contenido_bytes)
    elif nombre.endswith(".pdf"):
        texto = _extraer_pdf(contenido_bytes)
    else:
        raise ExtractError(
            "Formato no compatible. Sube un archivo .txt, .docx o .pdf, o pega el texto directamente."
        )

    texto = texto.strip()
    if not texto:
        raise ExtractError(
            "No se pudo leer texto en ese archivo (¿está vacío, o es una imagen escaneada sin texto?)."
        )
    if len(texto) > MAX_CHARS:
        texto = texto[:MAX_CHARS] + "\n\n[texto truncado por longitud]"
    return texto


def _decode_texto(data: bytes) -> str:
    for enc in ("utf-8", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    raise ExtractError("No se pudo leer la codificación del archivo de texto.")


def _extraer_docx(data: bytes) -> str:
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ExtractError("El archivo Word no se pudo abrir. ¿Es un .docx válido?") from exc
    partes = [p.text for p in doc.paragraphs if p.text.strip()]
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if celda.text.strip():
                    partes.append(celda.text.strip())
    return "\n".join(partes)


def _extraer_pdf(data: bytes) -> str:
    try:
        lector = pypdf.PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ExtractError("El archivo PDF no se pudo abrir. ¿Es un .pdf válido?") from exc
    if lector.is_encrypted:
        try:
            lector.decrypt("")
        except Exception:  # noqa: BLE001
            raise ExtractError("El PDF está protegido con contraseña; no se pudo leer.")
    partes = []
    inicio = time.monotonic()
    for i, pagina in enumerate(lector.pages):
        if i >= MAX_PDF_PAGINAS:
            partes.append("\n[se cortó la lectura por tener demasiadas páginas]")
            break
        if time.monotonic() - inicio > MAX_PDF_SEGUNDOS:
            partes.append("\n[se cortó la lectura por demorar demasiado]")
            break
        try:
            partes.append(pagina.extract_text() or "")
        except Exception:  # noqa: BLE001
            continue
    texto = "\n".join(partes)
    if not texto.strip():
        raise ExtractError(
            "El PDF parece ser una imagen escaneada sin texto seleccionable. "
            "Prueba subiendo la transcripción como .txt o .docx, o pégala directamente."
        )
    return texto
