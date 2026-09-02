"""
Genera el informe final descargable de un caso, en formato Word (.docx) —
reemplaza al informe en PDF: síntesis general, problemas identificados,
pasos según el reglamento interno (si se subió uno), sugerencias de acción,
el listado de relatos incluidos, y la firma del encargado.

Usa python-docx (ya es una dependencia del proyecto, para leer relatos en
Word) para no agregar ninguna librería nueva.
"""

import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x1E, 0x4E, 0x48)
INK = RGBColor(0x23, 0x20, 0x19)
MUTED = RGBColor(0x6E, 0x67, 0x5C)
URGENCIA_RGB = {
    "bajo": RGBColor(0x3F, 0x7D, 0x4F),
    "medio": RGBColor(0xA3, 0x61, 0x0C),
    "alto": RGBColor(0xB0, 0x36, 0x2A),
}


def _fmt_fecha(iso: str) -> str:
    try:
        dt = datetime.fromisoformat(iso)
        return dt.strftime("%d-%m-%Y %H:%M")
    except (ValueError, TypeError):
        return iso or "—"


def _titulo(doc, texto, tamano=18, color=INK):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(tamano)
    run.font.color.rgb = color
    return p


def _meta(doc, texto, color=MUTED):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(9.5)
    run.font.color.rgb = color
    return p


def _h2(doc, texto):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    run = p.add_run(texto.upper())
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = ACCENT
    return p


def _body(doc, texto):
    p = doc.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = INK
    return p


def _bullets(doc, items):
    items = items or ["—"]
    for it in items:
        p = doc.add_paragraph(str(it), style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = INK


def construir_informe_docx(caso, relatos, problemas, soluciones, pasos_reglamento=None, configuracion=None) -> bytes:
    """caso: sqlite3.Row de la tabla casos. relatos: lista de sqlite3.Row de la tabla relatos."""
    doc = Document()

    insignia_bytes = configuracion["insignia_bytes"] if configuracion else None
    if insignia_bytes:
        try:
            doc.add_picture(io.BytesIO(insignia_bytes), width=Inches(0.9))
        except Exception:
            pass  # imagen corrupta o formato no soportado por python-docx: se omite sin romper el informe

    _titulo(doc, "Informe de análisis — RelacionAI")
    nombre_colegio = configuracion["nombre_colegio"] if configuracion else None
    if nombre_colegio:
        _meta(doc, nombre_colegio, color=ACCENT)
    nivel = (caso["nivel_urgencia"] or "medio")
    _meta(doc, f"Caso {caso['rotulo']} · Apellido: {caso['apellido']}")
    linea_meta = doc.add_paragraph()
    run = linea_meta.add_run(
        f"Carpeta creada: {_fmt_fecha(caso['creado_en'])} · Relatos incluidos: {len(relatos)} · Nivel de urgencia: "
    )
    run.font.size = Pt(9.5)
    run.font.color.rgb = MUTED
    run_nivel = linea_meta.add_run(nivel.upper())
    run_nivel.bold = True
    run_nivel.font.size = Pt(9.5)
    run_nivel.font.color.rgb = URGENCIA_RGB.get(nivel, MUTED)
    if caso["sintesis_generada_en"]:
        _meta(doc, f"Síntesis generada: {_fmt_fecha(caso['sintesis_generada_en'])}")

    _h2(doc, "Síntesis general del caso")
    _body(doc, caso["sintesis_general"] or "Aún no se ha generado una síntesis para este caso.")

    _h2(doc, "Problemas identificados")
    _bullets(doc, problemas)

    if pasos_reglamento:
        _h2(doc, "Pasos del Reglamento Interno")
        _bullets(doc, pasos_reglamento)

    _h2(doc, "Sugerencias de acción")
    _bullets(doc, soluciones)

    _h2(doc, "Relatos incluidos en este caso")
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = "Light Grid Accent 1"
    hdr = tabla.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Persona", "Formato", "Recibido"
    for r in relatos:
        fila = tabla.add_row().cells
        fila[0].text = r["nombre_persona"]
        fila[1].text = r["formato_entrada"]
        fila[2].text = _fmt_fecha(r["subido_en"])

    doc.add_paragraph()
    nota = doc.add_paragraph()
    run = nota.add_run(
        "Generado con asistencia de Claude a partir de los relatos registrados en RelacionAI. "
        "Documento de uso interno y confidencial."
    )
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED

    nombre_encargado = configuracion["nombre_encargado"] if configuracion else None
    cargo_encargado = configuracion["cargo_encargado"] if configuracion else None
    if nombre_encargado or cargo_encargado:
        doc.add_paragraph()
        if nombre_encargado:
            p = doc.add_paragraph()
            run = p.add_run(nombre_encargado)
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = INK
        if cargo_encargado:
            _meta(doc, cargo_encargado)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def construir_relato_docx(caso, relato, numero=None, configuracion=None) -> bytes:
    """Un solo relato como documento Word descargable por sí solo — para que el encargado
    pueda guardar en su computador el respaldo de una persona en particular, sin tener que
    descargar el informe completo del caso. caso y relato: sqlite3.Row."""
    doc = Document()

    titulo = f"Relato — {relato['nombre_persona']}"
    if numero:
        titulo += f" (Relato {numero})"
    _titulo(doc, titulo)
    nombre_colegio = configuracion["nombre_colegio"] if configuracion else None
    if nombre_colegio:
        _meta(doc, nombre_colegio, color=ACCENT)
    _meta(doc, f"Caso {caso['rotulo']} · Apellido: {caso['apellido']}")
    _meta(doc, f"Recibido: {_fmt_fecha(relato['subido_en'])} · Formato: {relato['formato_entrada']}"
               + (f" · Archivo original: {relato['archivo_original']}" if relato["archivo_original"] else ""))

    if relato["resumen"]:
        _h2(doc, "Resumen")
        _body(doc, relato["resumen"])

    _h2(doc, "Relato completo")
    _body(doc, relato["contenido"] or "—")

    doc.add_paragraph()
    nota = doc.add_paragraph()
    run = nota.add_run(
        "Extraído de RelacionAI. Documento de uso interno y confidencial."
    )
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
